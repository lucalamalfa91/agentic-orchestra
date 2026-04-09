"""File-based knowledge source - reads local files and directories."""

import logging
from pathlib import Path
from typing import Optional

from .base_source import KnowledgeSource, Document

logger = logging.getLogger(__name__)


class FileSource(KnowledgeSource):
    """
    Reads local files and extracts content.

    Supports:
    - Plain text files (.txt, .md)
    - PDF files (.pdf) via pypdf
    - Word documents (.docx) via python-docx
    - Directories (recursively scans for supported files)
    """

    def __init__(
        self,
        paths: list[str],
        extensions: Optional[list[str]] = None,
        chunk_size: int = 500,
        overlap: int = 50,
    ):
        """
        Initialize the file source.

        Args:
            paths: List of file paths or directory paths
            extensions: File extensions to include (default: .txt, .md, .pdf, .docx)
            chunk_size: Approximate chunk size in words
            overlap: Word overlap between chunks
        """
        self.paths = [Path(p) for p in paths]
        self.extensions = extensions or [".txt", ".md", ".pdf", ".docx"]
        self.chunk_size = chunk_size
        self.overlap = overlap

    @property
    def name(self) -> str:
        # Use first path as identifier
        return f"files_{self.paths[0].name}"

    @property
    def description(self) -> str:
        return f"Local files from {', '.join(str(p) for p in self.paths)}"

    async def fetch(self, query: str) -> list[Document]:
        """
        Read all configured files and extract content.

        Args:
            query: Not used for file reading (all files processed)

        Returns:
            List of Document chunks
        """
        logger.info(f"[FileSource] Reading files from {len(self.paths)} paths")

        documents = []
        for path in self.paths:
            if path.is_file():
                docs = await self._process_file(path)
                documents.extend(docs)
            elif path.is_dir():
                docs = await self._process_directory(path)
                documents.extend(docs)
            else:
                logger.warning(f"[FileSource] Path not found: {path}")

        logger.info(f"[FileSource] Extracted {len(documents)} chunks from files")
        return documents

    async def _process_directory(self, directory: Path) -> list[Document]:
        """Recursively process all files in a directory."""
        documents = []
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix in self.extensions:
                docs = await self._process_file(file_path)
                documents.extend(docs)
        return documents

    async def _process_file(self, file_path: Path) -> list[Document]:
        """Process a single file and return chunked documents."""
        try:
            extension = file_path.suffix.lower()

            if extension in [".txt", ".md"]:
                content = await self._read_text_file(file_path)
            elif extension == ".pdf":
                content = await self._read_pdf_file(file_path)
            elif extension == ".docx":
                content = await self._read_docx_file(file_path)
            else:
                logger.warning(f"[FileSource] Unsupported file type: {extension}")
                return []

            # Chunk content
            chunks = self._chunk_content(content)
            documents = [
                Document(
                    content=chunk,
                    source=str(file_path),
                    metadata={
                        "filename": file_path.name,
                        "extension": extension,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                    },
                )
                for i, chunk in enumerate(chunks)
            ]

            logger.debug(f"[FileSource] Processed {file_path.name}: {len(documents)} chunks")
            return documents

        except Exception as e:
            logger.error(f"[FileSource] Error processing {file_path}: {e}")
            return []

    async def _read_text_file(self, file_path: Path) -> str:
        """Read plain text or markdown file."""
        return file_path.read_text(encoding="utf-8")

    async def _read_pdf_file(self, file_path: Path) -> str:
        """Read PDF file using pypdf."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("[FileSource] pypdf not installed - cannot read PDF files")
            return ""

    async def _read_docx_file(self, file_path: Path) -> str:
        """Read Word document using python-docx."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)

        except ImportError:
            logger.error("[FileSource] python-docx not installed - cannot read DOCX files")
            return ""

    def _chunk_content(self, content: str) -> list[str]:
        """Split content into chunks with overlap."""
        if not content:
            return []

        words = content.split()
        chunks = []

        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i : i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            if chunk_text.strip():
                chunks.append(chunk_text)

        return chunks
